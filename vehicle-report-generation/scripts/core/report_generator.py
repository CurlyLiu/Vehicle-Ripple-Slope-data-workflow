"""报告生成器基类."""

import shutil

from docx import Document
from docx.shared import Inches, Pt

from ..utils.condition_mapper import (
    REPORT_ITEMS,
    count_matched_images,
    count_matched_items,
    find_image_records,
    find_report_item_data,
)
from ..utils.docx_helpers import (
    insert_table_at_top,
    remove_section_by_heading,
    remove_table_rows,
)
from ..utils.image_resolver import resolve_image_path

SOC_LEVELS = ["≥70%", "40%-70%", "≤40%"]
TOTAL_ITEMS_PER_SOC = 9
TOTAL_IMAGES_PER_SOC = 16
EXPECTED_IMAGE_TABLE_ROWS = 1 + TOTAL_IMAGES_PER_SOC * 2  # 1 header + 16 image/caption pairs

NO_DATA_TEXT = "未找到对应工况数据。"


class ReportGenerator:
    """报告生成器基类."""

    report_type_label: str = ""

    def __init__(self, template_path: str, prune: bool = True):
        self.template_path = template_path
        self.prune = prune

    def load_data(self, vehicle_id: str, component_code: str, soc_level: str) -> list[dict]:
        """子类必须实现：加载指定车辆、组件、SOC区间的数据记录."""
        raise NotImplementedError

    def build_result_text(self, item_index: int, matched_data: dict) -> str:
        """子类必须实现：根据匹配数据生成检验结果文本."""
        raise NotImplementedError

    def build_compliance(self, item_index: int, matched_data: dict) -> str:
        """子类必须实现：根据匹配数据生符合性判定文本."""
        raise NotImplementedError

    def adapt_standard_requirement(self, original_text: str) -> str:
        """子类可重写：修改标准要求列文本（斜率报告需去掉限值）."""
        return original_text

    def _rewrite_paragraph_text(self, paragraph) -> None:
        """跨 run 拼接重建段落文本 (处理 python-docx run 拆分陷阱).

        同一可见标题可能被拆成多个 Run 对象 (粗体/颜色变化处自动分 run)。
        逐 run replace 会漏掉跨 run 的子串。

        此方法:
        1. 用 paragraph.text 跨 run 拼接全文本
        2. 应用 adapt_standard_requirement
        3. 若长度未变 (如"电压纹波→电流纹波"等长替换),按 run 边界回写以保留各 run 的字符样式 (REPORT-H1)
        4. 若长度变化 (如"30Vpp→100App"),折叠到第一个 run + 清空其余 run.text (会丢失原 run1+ 的样式,但功能正确)
        5. 幂等: 无变化时跳过
        """
        full_text = paragraph.text
        if not full_text.strip():
            return
        new_text = self.adapt_standard_requirement(full_text)
        if new_text == full_text:
            return  # 幂等: 无变化跳过

        # REPORT-H1 v1.4: 等长替换按 run 边界回写,保留字符样式 (粗体/颜色/字号)
        if len(new_text) == len(full_text) and paragraph.runs:
            offset = 0
            for run in paragraph.runs:
                run_len = len(run.text)
                run.text = new_text[offset:offset + run_len]
                offset += run_len
            return

        # 长度变化: 折叠到第一个 run (会丢失 run1+ 的字符样式,可接受 - 整段统一格式)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new_text)

    def _rewrite_titles_and_headers(self, doc) -> None:
        """改写章节标题与表格首行表头 (电流/电压通道单位适配,幂等).

        v1.4 P1.3: 仅 _fill_result_table 改写表格 cell,章节标题段落与表头
        字面含"电压纹波/30Vpp/电压斜率",此方法补全这两类位置。
        """
        # 章节标题段落
        for paragraph in doc.paragraphs:
            if paragraph.style.name in ("Heading 1", "Heading 2", "Heading 3"):
                self._rewrite_paragraph_text(paragraph)
        # 表格首行 (表头)
        for table in doc.tables:
            if not table.rows:
                continue
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    self._rewrite_paragraph_text(paragraph)

    def generate(self, vehicle_id: str, component_code: str, output_path: str) -> None:
        """生成报告."""
        # 复制模板
        shutil.copy(self.template_path, output_path)

        doc = Document(output_path)

        # P1.3: 改写章节标题与表头 (电流/电压通道适配)
        self._rewrite_titles_and_headers(doc)

        empty_sections: list[str] = []
        soc_stats: dict[str, dict[str, int]] = {}

        for soc_index, soc_level in enumerate(SOC_LEVELS):
            data = self.load_data(vehicle_id, component_code, soc_level)

            # 检验结果表格（偶数索引：0, 2, 4）
            result_table = doc.tables[soc_index * 2]
            data_rows, empty_row_indices = self._fill_result_table(result_table, data)

            # 试验数据曲线表格（奇数索引：1, 3, 5）
            image_table = doc.tables[soc_index * 2 + 1]
            img_count, empty_image_row_indices = self._fill_image_table(image_table, data)

            soc_stats[soc_level] = {"items": data_rows, "images": img_count}

            if self.prune:
                if data_rows == 0 and img_count == 0:
                    # 整章删除（在所有 SOC 处理完后统一删除）
                    empty_sections.append(soc_level)
                else:
                    # 行级与图片对级裁剪
                    remove_table_rows(result_table, empty_row_indices)
                    remove_table_rows(image_table, empty_image_row_indices)

                    # 裁剪后重新编号
                    self._renumber_result_table(result_table)
                    self._renumber_image_table(image_table)

        # 章节级裁剪（在统计完成后再删，避免影响 doc.tables 索引）
        if self.prune:
            for soc_level in empty_sections:
                remove_section_by_heading(doc, soc_level, next_headings=SOC_LEVELS)

        # 插入测试覆盖度摘要表
        if self.prune:
            self._insert_coverage_summary(
                doc,
                vehicle_id=vehicle_id,
                component_code=component_code,
                soc_stats=soc_stats,
            )

            # 极端情况：所有 SOC 都被删除，加一段提示
            total_items = sum(s["items"] for s in soc_stats.values())
            total_images = sum(s["images"] for s in soc_stats.values())
            if total_items == 0 and total_images == 0:
                self._append_empty_data_notice(doc)

        doc.save(output_path)

    def _fill_result_table(self, table, records: list[dict]) -> tuple[int, list[int]]:
        """填充检验结果表格，返回 (有数据行数, 空行索引列表).

        空行索引基于原始表格 (1..9 对应 9 个工况)。
        """
        empty_rows: list[int] = []
        data_rows = 0

        for item_index in range(TOTAL_ITEMS_PER_SOC):
            row_index = item_index + 1
            row = table.rows[row_index]
            matched = find_report_item_data(item_index, records)

            # 标准要求列（第4列，索引3）- 斜率报告需适配
            std_cell = row.cells[3]
            original_text = std_cell.text
            adapted_text = self.adapt_standard_requirement(original_text)
            if adapted_text != original_text:
                for p in std_cell.paragraphs:
                    p.clear()
                std_cell.paragraphs[0].add_run(adapted_text)

            # 检验结果列（第5列，索引4）
            result_cell = row.cells[4]
            result_text = self.build_result_text(item_index, matched)
            for p in result_cell.paragraphs:
                p.clear()
            result_cell.paragraphs[0].add_run(result_text)

            # 符合性判定列（第6列，索引5）
            compliance_cell = row.cells[5]
            compliance_text = self.build_compliance(item_index, matched)
            for p in compliance_cell.paragraphs:
                p.clear()
            compliance_cell.paragraphs[0].add_run(compliance_text)

            if result_text == NO_DATA_TEXT:
                empty_rows.append(row_index)
            else:
                data_rows += 1

        return data_rows, empty_rows

    def _fill_image_table(self, table, records: list[dict]) -> tuple[int, list[int]]:
        """填充试验数据曲线表格，返回 (有图数, 待删除行索引列表).

        每张图占 2 行（图行 + 图注行）。如果模板行数与预期不符，
        跳过图片级裁剪以保持向后兼容。
        """
        image_records = find_image_records(records)
        empty_rows: list[int] = []
        img_count = 0

        # 模板防御：行数应为 1 + N*2（N=16）
        template_ok = len(table.rows) == EXPECTED_IMAGE_TABLE_ROWS

        for img_index, rec in enumerate(image_records):
            empty_row_index = 1 + img_index * 2
            if empty_row_index >= len(table.rows):
                continue

            cell = table.rows[empty_row_index].cells[0]
            for para in cell.paragraphs:
                para.clear()

            img_path = None
            if rec and rec.get("image_path"):
                img_path = resolve_image_path(rec["image_path"])

            if img_path:
                run = cell.paragraphs[0].add_run()
                run.add_picture(img_path, width=Inches(5.5))
                img_count += 1
            elif template_ok:
                # 标记图行 + 图注行待删
                empty_rows.append(empty_row_index)
                caption_row_index = 2 + img_index * 2
                if caption_row_index < len(table.rows):
                    empty_rows.append(caption_row_index)

        return img_count, empty_rows

    def _insert_coverage_summary(
        self,
        doc,
        vehicle_id: str,
        component_code: str,
        soc_stats: dict[str, dict[str, int]],
    ) -> None:
        """在文档开头插入测试覆盖度摘要表."""
        tested_socs = [s for s, st in soc_stats.items() if st["items"] > 0 or st["images"] > 0]
        total_items_actual = sum(s["items"] for s in soc_stats.values())
        total_images_actual = sum(s["images"] for s in soc_stats.values())

        total_items_max = TOTAL_ITEMS_PER_SOC * len(SOC_LEVELS)
        total_images_max = TOTAL_IMAGES_PER_SOC * len(SOC_LEVELS)

        rating = self._format_completeness_rating(
            tested_socs, total_items_actual, total_items_max,
            total_images_actual, total_images_max,
        )

        rows = [
            ("车辆编号", vehicle_id),
            ("组件通道", component_code),
            ("报告类型", self.report_type_label or "—"),
            (
                "已测 SOC 区间",
                f"{', '.join(tested_socs)}（共 {len(tested_socs)}/{len(SOC_LEVELS)}）"
                if tested_socs
                else f"无（0/{len(SOC_LEVELS)}）",
            ),
            (
                "已测工况数",
                f"{total_items_actual} / {total_items_max}",
            ),
            (
                "已测试验曲线图数",
                f"{total_images_actual} / {total_images_max}",
            ),
            ("数据完整度", rating),
        ]

        table = insert_table_at_top(doc, rows=len(rows) + 1, cols=2)
        try:
            table.style = "Table Grid"
        except KeyError:
            # 某些环境下缺少该样式，跳过即可
            pass

        # 表头
        header = table.rows[0].cells
        header[0].text = "项目"
        header[1].text = "数值/状态"
        for cell in header:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        # 数据行
        for i, (label, value) in enumerate(rows, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value

    def _format_completeness_rating(
        self,
        tested_socs: list[str],
        items_actual: int,
        items_max: int,
        images_actual: int,
        images_max: int,
    ) -> str:
        """根据覆盖度计算完整度评级文字."""
        if items_actual == 0 and images_actual == 0:
            return "无数据"
        if (
            len(tested_socs) == len(SOC_LEVELS)
            and items_actual == items_max
            and images_actual == images_max
        ):
            return "完整覆盖"
        return "部分覆盖"

    def _renumber_result_table(self, table) -> None:
        """对检验结果表的数据行重新编号序号列（第0列），保留原始格式."""
        for i, row in enumerate(table.rows):
            if i == 0:
                continue
            seq_cell = row.cells[0]
            for p in seq_cell.paragraphs:
                for run in p.runs:
                    text = run.text.strip()
                    if text.isdigit():
                        run.text = str(i)
                        break

    def _renumber_image_table(self, table) -> None:
        """对图片表的图注行重新编号（替换开头的'图X'），保留原始格式."""
        import re

        for i, row in enumerate(table.rows):
            if i == 0:
                continue
            if i % 2 == 0:
                new_num = i // 2
                caption_cell = row.cells[0]
                for p in caption_cell.paragraphs:
                    runs = list(p.runs)
                    for idx, run in enumerate(runs):
                        # 情况1：前一run含"图"，当前run是纯数字
                        if (
                            idx > 0
                            and "图" in runs[idx - 1].text
                            and run.text.strip().isdigit()
                        ):
                            run.text = str(new_num)
                            break
                        # 情况2：当前run以"图"开头且后面紧跟数字
                        match = re.match(r"^(图\s*)(\d+)(.*)", run.text)
                        if match:
                            run.text = f"{match.group(1)}{new_num}{match.group(3)}"
                            break

    def _append_empty_data_notice(self, doc) -> None:
        """在文档末尾(sectPr 之前)插入一段提示文字."""
        para = doc.add_paragraph("该组件未采集到任何有效数据。")
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(12)
