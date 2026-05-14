"""Tests for docx_helpers XML manipulation utilities."""

import sys
from pathlib import Path

import pytest
from docx import Document

try:
    from docx.oxml import qn
except ImportError:
    from docx.oxml.ns import qn

# Ensure scripts are importable
sys.path.insert(0, str(Path(__file__).parent.parent))

from scripts.utils.docx_helpers import (
    insert_table_at_top,
    remove_section_by_heading,
    remove_table_rows,
)


class TestRemoveTableRows:
    def test_remove_single_row(self):
        doc = Document()
        table = doc.add_table(rows=5, cols=2)
        for i, row in enumerate(table.rows):
            row.cells[0].text = f"A{i}"
            row.cells[1].text = f"B{i}"

        remove_table_rows(table, [2])
        assert len(table.rows) == 4
        assert table.rows[2].cells[0].text == "A3"

    def test_remove_multiple_rows_descending(self):
        doc = Document()
        table = doc.add_table(rows=5, cols=2)
        for i, row in enumerate(table.rows):
            row.cells[0].text = f"A{i}"

        remove_table_rows(table, [1, 3])
        assert len(table.rows) == 3
        texts = [row.cells[0].text for row in table.rows]
        assert texts == ["A0", "A2", "A4"]

    def test_remove_all_rows(self):
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        for i, row in enumerate(table.rows):
            row.cells[0].text = f"A{i}"

        remove_table_rows(table, [0, 1, 2])
        assert len(table.rows) == 0

    def test_remove_out_of_range_ignored(self):
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        remove_table_rows(table, [10])
        assert len(table.rows) == 3


class TestRemoveSectionByHeading:
    def test_remove_middle_section(self):
        doc = Document()
        doc.add_paragraph("前言")
        doc.add_paragraph("≥70%")
        doc.add_table(rows=2, cols=2)
        doc.add_paragraph("40%-70%")
        doc.add_table(rows=2, cols=2)

        assert remove_section_by_heading(doc, "≥70%", next_headings=["≥70%", "40%-70%", "≤40%"])

        body = list(doc.element.body)
        texts = []
        tables = 0
        for elem in body:
            if elem.tag.endswith("}p"):
                t = "".join(t.text or "" for t in elem.findall(f".//{qn('w:t')}"))
                texts.append(t)
            elif elem.tag.endswith("}tbl"):
                tables += 1

        assert "前言" in texts
        assert "≥70%" not in texts
        assert "40%-70%" in texts
        assert tables == 1

    def test_remove_last_section_protects_sectpr(self):
        doc = Document()
        doc.add_paragraph("≥70%")
        doc.add_table(rows=2, cols=2)
        doc.add_paragraph("≤40%")
        doc.add_table(rows=2, cols=2)

        assert remove_section_by_heading(doc, "≤40%", next_headings=["≥70%", "40%-70%", "≤40%"])

        body = list(doc.element.body)
        has_sectpr = any(elem.tag.endswith("}sectPr") for elem in body)
        assert has_sectpr

        texts = []
        for elem in body:
            if elem.tag.endswith("}p"):
                t = "".join(t.text or "" for t in elem.findall(f".//{qn('w:t')}"))
                texts.append(t)

        assert "≥70%" in texts
        assert "≤40%" not in texts

    def test_heading_not_found(self):
        doc = Document()
        doc.add_paragraph("前言")
        assert not remove_section_by_heading(doc, "不存在")


class TestInsertTableAtTop:
    def test_insert_table_at_top_existing_paragraphs(self):
        doc = Document()
        doc.add_paragraph("第一段")
        doc.add_paragraph("第二段")

        table = insert_table_at_top(doc, rows=2, cols=2)

        body = list(doc.element.body)
        assert body[0] is table._element

        # sectPr should still be at the end
        assert body[-1].tag.endswith("}sectPr")

    def test_insert_table_at_top_empty_doc(self):
        doc = Document()
        # Empty doc only has sectPr
        table = insert_table_at_top(doc, rows=2, cols=2)

        body = list(doc.element.body)
        assert body[0] is table._element
        assert body[-1].tag.endswith("}sectPr")
